#!/usr/bin/env python3
"""Assembler for the Daily Intel Brief.

Composes the analyst's regional JSON + the visuals manifest into the
structured JSON shape that skills/pdf-report/scripts/render_pdf.py
expects, then invokes the pdf-report renderer to produce the PDF.

Also writes a parallel .docx via python-docx so the principal can
revise the working draft.

Usage:

    python3 scripts/build_pdf.py --working-dir <wd> \
        --out-pdf <wd>/final/brief-YYYY-MM-DD.pdf \
        --out-docx <wd>/final/brief-YYYY-MM-DD.docx
"""
from __future__ import annotations

import argparse
import datetime as dt
import json
import os
import pathlib
import shutil
import subprocess
import sys
from collections import Counter
from typing import Any

REPO_ROOT_GUESS = pathlib.Path(__file__).resolve().parents[3]
PDF_REPORT_RENDERER = REPO_ROOT_GUESS / "skills" / "pdf-report" / "scripts" / "render_pdf.py"
PDF_VENV = pathlib.Path("~/.openclaw/workspace/.venv_pdf/bin/python").expanduser()

REGIONS_ORDER = [
    "europe", "asia", "middle_east",
    "north_america", "south_central_america", "global_finance",
]
REGION_LABEL = {
    "europe": "Europe",
    "asia": "Asia",
    "middle_east": "Middle East",
    "north_america": "North America",
    "south_central_america": "South & Central America (incl. Caribbean)",
    "global_finance": "Global Finance",
}


def log(msg: str) -> None:
    ts = dt.datetime.utcnow().strftime("%H:%M:%S")
    print(f"[assemble {ts}] {msg}", file=sys.stderr, flush=True)


def load_json(p: pathlib.Path) -> Any:
    return json.loads(p.read_text())


def find_visual(manifest: dict, region: str, kind: str) -> str | None:
    for a in manifest.get("assets", []):
        if a.get("region") == region and a.get("kind") == kind:
            return a.get("path")
    return None


def incidents_for_region(incidents: list[dict], region: str) -> list[dict]:
    return [i for i in incidents if i.get("region") == region]


def build_pdfreport_payload(wd: pathlib.Path) -> dict:
    incidents_payload = load_json(wd / "raw" / "incidents.json")
    incidents = incidents_payload.get("incidents", [])
    exec_summary = load_json(wd / "analysis" / "exec_summary.json")
    visuals_manifest_path = wd / "visuals" / "manifest.json"
    manifest = load_json(visuals_manifest_path) if visuals_manifest_path.exists() else {"assets": []}
    red_team_path = wd / "analysis" / "red_team.md"
    red_team = red_team_path.read_text() if red_team_path.exists() else ""

    today = dt.datetime.utcnow().strftime("%Y-%m-%d")
    long_date = dt.datetime.utcnow().strftime("%A %d %B %Y")
    dtg = dt.datetime.utcnow().strftime("%Y%m%dT%H%MZ")

    sections: list[dict] = []

    # Cover + exec are baked into the title/summary fields below;
    # the per-section list starts at the regional sections.
    for region in REGIONS_ORDER:
        region_path = wd / "analysis" / f"{region}.json"
        if not region_path.exists():
            continue
        rp = load_json(region_path)
        region_label = REGION_LABEL[region]
        region_incidents = incidents_for_region(incidents, region)

        # Items: factual narrative + key judgments + indicator block
        items = []
        items.append(rp.get("narrative", ""))
        for kj in rp.get("key_judgments", []) or []:
            items.append(
                f"{kj.get('id')} ({kj.get('sherman_kent_band')}; "
                f"{kj.get('prediction_pct')}% / {kj.get('horizon_days')}d): "
                f"{kj.get('statement')}"
            )
            ind = kj.get("what_would_change_it") or []
            if ind:
                items.append("  Indicators: " + " // ".join(ind))

        # Tabular incident list
        rows = []
        for inc in region_incidents:
            srcs = inc.get("sources") or []
            srcrating = ""
            if srcs:
                s = srcs[0]
                srcrating = f"{s.get('admiralty_reliability','?')}{s.get('admiralty_credibility','?')}"
            rows.append([
                inc["id"][-4:],
                (inc.get("country") or "—"),
                inc.get("category", ""),
                (inc.get("headline") or "")[:90],
                srcrating,
            ])

        # Visual
        chart_kind = "finance_panel" if region == "global_finance" else "regional_map"
        chart_path = find_visual(manifest, region, chart_kind)
        charts = []
        if chart_path:
            charts.append({
                "title": region_label,
                "src": chart_path,
                "caption": (
                    f"{region_label} — {len(region_incidents)} "
                    f"{'market events' if region == 'global_finance' else 'incidents'} "
                    f"in 24h to {dtg}."
                ),
            })
        # Relationships, if it exists for this region
        rel_path = find_visual(manifest, region, "relationships")
        if rel_path:
            charts.append({
                "title": f"{region_label} — actor / event graph",
                "src": rel_path,
                "caption": "Solid edges = observed; dashed edges = claimed.",
            })

        sections.append({
            "title": region_label,
            "lead": f"{len(region_incidents)} incidents in 24h to {dtg} UTC",
            "items": items,
            "table": {
                "headers": ["Pin", "Country", "Category", "Headline", "Src"],
                "rows": rows,
            } if rows else None,
            "charts": charts,
            "note": None,
        })

    # Annex A — sources + collection gaps
    sources_count = Counter()
    for inc in incidents:
        for s in inc.get("sources") or []:
            sources_count[s.get("name", "?")] += 1
    annex_a_items = [
        f"{name} ({count})"
        for name, count in sources_count.most_common()
    ]
    if incidents_payload.get("collection_gaps"):
        annex_a_items.append("")
        annex_a_items.append("Collection gaps in this window:")
        annex_a_items.extend(f"- {g}" for g in incidents_payload["collection_gaps"])
    sections.append({
        "title": "Annex A — Sources & Methodology",
        "lead": (
            "Source ratings: NATO Admiralty Code (A1 ... F6) per "
            "skills/source-evaluation. Confidence bands: Sherman Kent. "
            "Predictions are 7-day, falsifiable, scored at the monthly "
            "calibration review."
        ),
        "items": annex_a_items,
        "table": None,
        "charts": [],
        "note": None,
    })

    # Annex B — I&W status board (placeholder; fed by analyst/iw-boards if present)
    iw_boards_dir = REPO_ROOT_GUESS / "analyst" / "iw-boards"
    iw_items = []
    if iw_boards_dir.exists():
        for f in sorted(iw_boards_dir.glob("*.md")):
            iw_items.append(f"See {f.relative_to(REPO_ROOT_GUESS)} for standing board.")
    if not iw_items:
        iw_items = ["No standing I&W boards present. Stand one up the first time a topic merits two consecutive daily appearances."]
    sections.append({
        "title": "Annex B — Indicators & Warnings — Status Board",
        "lead": "Standing indicators that fired or moved in the last 24h.",
        "items": iw_items,
        "table": None,
        "charts": [],
        "note": None,
    })

    # Annex C — red team
    if red_team.strip():
        sections.append({
            "title": "Annex C — Red-team Note",
            "lead": "Steel-manned alternative to today's most load-bearing key judgment.",
            "items": [red_team],
            "table": None,
            "charts": [],
            "note": None,
        })

    # Cover-level summary list (top 5 from exec_summary)
    summary = []
    summary.append(f"BLUF: {exec_summary.get('bluf','')}")
    summary.append(exec_summary.get("context_paragraph", ""))
    for kj in exec_summary.get("five_judgments", []):
        summary.append(
            f"[{kj.get('drawn_from_region')}] {kj.get('statement')} "
            f"({kj.get('sherman_kent_band')}; {kj.get('prediction_pct')}%/7d)"
        )

    payload = {
        "title": "TREVOR DAILY INTELLIGENCE BRIEF",
        "subtitle": f"{long_date} — DTG {dtg}",
        "generated_at": dt.datetime.utcnow().strftime("%Y-%m-%d %H:%M Z"),
        "summary": summary,
        "sections": sections,
        "footer": "TREVOR Daily — UNCLASSIFIED // FOR OFFICIAL USE — Methodology: NATO Admiralty + Sherman Kent",
    }
    return payload


def write_docx(payload: dict, out_path: pathlib.Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        log("python-docx not installed; skipping .docx (pip install python-docx)")
        return False
    doc = Document()
    doc.add_heading(payload["title"], level=0)
    doc.add_paragraph(payload["subtitle"])
    doc.add_paragraph(payload["generated_at"])
    doc.add_heading("Executive Summary", level=1)
    for line in payload["summary"]:
        if line:
            doc.add_paragraph(line)
    for section in payload["sections"]:
        doc.add_heading(section["title"], level=1)
        if section.get("lead"):
            p = doc.add_paragraph()
            run = p.add_run(section["lead"]); run.italic = True
        for item in section.get("items", []):
            if item:
                doc.add_paragraph(item)
        if section.get("table"):
            t = section["table"]
            tbl = doc.add_table(rows=1, cols=len(t["headers"]))
            tbl.style = "Light List Accent 1"
            hdr = tbl.rows[0].cells
            for i, h in enumerate(t["headers"]):
                hdr[i].text = h
            for row in t["rows"]:
                cells = tbl.add_row().cells
                for i, v in enumerate(row):
                    cells[i].text = str(v)
        for ch in section.get("charts", []):
            try:
                src = ch["src"]
                if not os.path.isabs(src):
                    src = str((out_path.parent.parent / src).resolve())
                if os.path.exists(src):
                    doc.add_picture(src)
                    if ch.get("caption"):
                        cap = doc.add_paragraph(ch["caption"])
                        cap.runs[0].italic = True
            except Exception as exc:
                log(f"docx image embed failed for {ch.get('src')}: {exc}")
    doc.add_paragraph()
    doc.add_paragraph(payload.get("footer", ""))
    doc.save(str(out_path))
    return True


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--working-dir", required=True)
    parser.add_argument("--out-pdf", required=True)
    parser.add_argument("--out-docx", default=None)
    parser.add_argument("--template",
                        default=str(pathlib.Path(__file__).resolve().parent.parent
                                    / "templates" / "daily-product.html"))
    args = parser.parse_args()

    wd = pathlib.Path(args.working_dir).expanduser().resolve()
    out_pdf = pathlib.Path(args.out_pdf).expanduser().resolve()
    out_pdf.parent.mkdir(parents=True, exist_ok=True)

    payload = build_pdfreport_payload(wd)

    pdfdata_path = wd / "final" / "report-data.json"
    pdfdata_path.write_text(json.dumps(payload, indent=2))

    # Stage the template inside the working dir so the pdf-report
    # workspace check passes. (skills/pdf-report enforces that --input,
    # --output, --template-file, and chart src paths all sit inside the
    # OPENCLAW workspace; the simplest way to satisfy that is to point
    # the workspace at $wd and put the template there too.)
    staged_template = wd / "templates" / "daily-product.html"
    staged_template.parent.mkdir(parents=True, exist_ok=True)
    src_template = pathlib.Path(args.template)
    if src_template.exists():
        shutil.copy(src_template, staged_template)
    else:
        staged_template = None

    # Invoke the pdf-report skill renderer
    interp = str(PDF_VENV) if PDF_VENV.exists() else sys.executable
    if not PDF_REPORT_RENDERER.exists():
        log(f"WARN: pdf-report renderer not found at {PDF_REPORT_RENDERER}; "
            "wrote payload JSON only")
        return 0
    cmd = [
        interp, str(PDF_REPORT_RENDERER),
        "--input", str(pdfdata_path),
        "--output", str(out_pdf),
    ]
    if staged_template:
        cmd.extend(["--template-file", str(staged_template)])
    env = dict(os.environ)
    env["OPENCLAW_WORKSPACE"] = str(wd)
    log(f"invoking pdf-report renderer: {' '.join(cmd)}")
    rc = subprocess.call(cmd, env=env)
    if rc != 0:
        log(f"pdf-report renderer rc={rc}")
        return rc

    if args.out_docx:
        write_docx(payload, pathlib.Path(args.out_docx).expanduser().resolve())
    return 0


if __name__ == "__main__":
    sys.exit(main())
